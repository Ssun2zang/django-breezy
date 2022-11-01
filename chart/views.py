import re
from django.shortcuts import render
import pymysql
from datetime import datetime, timedelta
import pprint
from influxdb import InfluxDBClient
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import numpy as np
from scipy import signal
import math
import matplotlib
matplotlib.use('Qt5agg')
import matplotlib.pyplot as plt
from threading import Thread
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches
from django.http import Http404
from django.core.paginator import Paginator


def get_ifdb(db, host='localhost', port=8086):
    
    client = InfluxDBClient(host, port, database = db)
    try:
        client.create_database(db)
        print('Connection Successful')
        print(f'host :{host}\nport :{port}\n\ndatabase :{db}')
    except:
        print('Connection Failed')
        pass
    return client

def printresult(ifdb, tag):
    result = ifdb.query('SELECT "acc_1","acc_2","acc_3", "str_1", "str_2", "str_3" \
        FROM JANET_1 WHERE "S_id" = \''+tag+'\'')
    return result
    # pprint.pprint(result.raw)
    
def do_test(tag):
    mydb = get_ifdb(db='sensor')
    # print(mydb.get_list_database())
    _list = printresult(mydb, tag)
    return _list

# Create your views here.

def makereport(_document, _list, BRID):
    document = Document(_document)
    picname = 'static/totalpicture'+BRID+'.png'
    document.add_picture(picname,width= Cm(16), height= Cm(9))

    table = document.tables[0]  # 표가 1개만 있으니 0번째 표를 대입
    n_add_rows = len(_list) + 1 - len(table.rows) # 더 만들어야 할 행 개수


    for i in range(n_add_rows):  # 행 추가
        table.add_row()
    for i in range(n_add_rows+1):  # result 값을 표에 입력
        for j in range(4):
            table.cell(i+1,j).text = _list[i][j]
            table.rows[i+1].cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # 표에 들어가는 타입은 문자열, 따라서 정수 및 실수는 문자열로 변환
    # table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_name = "static/output"+BRID+".docx"
    document.save(report_name) # 다른 이름으로 저장


def makefgraph(acczlist, time):    ###### 시간을 못찾음
    b = np.array(acczlist)
    b = np.reshape(b, (len(b),1))
    # print(b)
    # print(type(b))
    fs=100
    f, Pacc = signal.csd(np.transpose(b),np.transpose(b), fs, nperseg=2048,noverlap=1024)

    #Pacc.shape
    f=f.reshape(-1,1)
    # b, a = signal.butter(4, 1/50, 'hp')
    # Pacc = signal.filtfilt(b, a, Pacc, padlen=50)
    
    Pacc=np.transpose(Pacc)
    
    print(float(abs(max(Pacc))), "최대값")
    print(time[int(str(np.where(Pacc == max(Pacc))[0])[1:-1])][:19])
    stamp = time[int(str(np.where(Pacc == max(Pacc))[0])[1:-1])][:19]
    new_stamp = stamp.replace('T', '_')
    print(float(f[np.where(Pacc == max(Pacc))]), "최대값일 때 주파수값")   # spot 2


    print(math.sqrt(float(abs(max(Pacc)))))  # spot 1

    spot1 = float(f[np.where(Pacc == max(Pacc))])  # x축
    spot2 = math.sqrt(float(abs(max(Pacc))))  # y축
    

    f = open('chart/spot.txt', 'a')
    stringgg = new_stamp+ " " + str(round(spot1,2)) + " " + str(round(spot2,2)) + "\n"
    f.write(stringgg)
    print(stringgg)
    f.close()



def totalgraph(_list1, _list2, BRID):
    from matplotlib import font_manager, rc
    font_path = "C:/Windows/Fonts/NGULIM.TTF"
    font = font_manager.FontProperties(fname=font_path).get_name()
    rc('font', family=font)

    data = np.loadtxt('chart/accgraph.txt')

    x=data[:,0]
    a1=data[:,1]
    a2=data[:,2]
    a3=data[:,3]


    # print(x)
    spots = []
    f = open('chart/spot.txt', 'r')
    for line in f.readlines():
            spot = line.split(' ')
            spots.append(spot)
    f.close()

    # print(spots[0][1][:-2])
    # print(type(spots[0][1]))

    for i in range(0,len(spots)):
        t = spots[i][0]
        a = float(spots[i][1])
        # print(a)
        b = float(spots[i][2][:-1])   # 피크
        # print(b)
        _a = round(a, 1)    # 주파수
        compa = np.where(x == _a)[0]
        # print(compa)

        result = "ERR"
        if b >= a3[compa]:
            result = "D"
        elif b >= a2[compa]:
            result = "C"
        elif b >= a1[compa]:
            result = "B"
        else:
            result = "A"

        
        print(str(a),str(b),"위험 수준 : "+result+"\n")
        if (b <= 0.01):
            b = 0.01

        _list2.append([t, a, b, result])

        plt.scatter(a,b,c="red")

    x=data[:,0].tolist()
    a1=data[:,1].tolist()
    a2=data[:,2].tolist()
    a3=data[:,3].tolist()
    _list1.append(x)
    _list1.append(a1)
    _list1.append(a2)
    _list1.append(a3)

    plt.plot(x, a1, label = 'A', color='yellowgreen')
    plt.plot(x, a2, label = 'B', color='darkorange')
    plt.plot(x, a3, label = 'C', color='darkred')
    plt.text(4.9, 0.1, 'A', fontsize=20, color='yellowgreen')
    plt.text(2,  0.55, 'B', fontsize=20, color='gold')
    plt.text(1.55, 1.4, 'C', fontsize=20, color='orange')
    plt.text(0.7, 10, 'D', fontsize=20, color='red')
    # plt.legend(loc='best')

    plt.xlabel('주파수(Hz)', fontsize=15, weight='bold')
    plt.ylabel('spectral power', fontsize=15, weight='bold')
    plt.semilogx()
    plt.semilogy()

    picname = "static/totalpicture"+BRID+".png"

    plt.savefig(picname)
    # plt.show()
    return



# def loading(request, BRID, year, month, day, year2, month2, day2):


def datedata(request, BRID):

    userform = request.session.get('user')

    conn = pymysql.connect(
                user='root',
                passwd='pass',
                host='52.79.181.152',
                port=59759,
                db='breezy'
            )

    curs = conn.cursor(pymysql.cursors.DictCursor)
            
    sql = """update `breezy`.`report` set `cur` = %s where (`rp_sID` = %s and `writer` = %s)"""

    curs.execute(sql, ('0', BRID, userform))
    conn.commit()

    return render(request, 'wait.html')

def _datedata(request, BRID, year, month, day, year2, month2, day2):

    userform = request.session.get('user')
    Success = 0
    conn = pymysql.connect(
                    user='root',
                    passwd='pass',
                    host='52.79.181.152',
                    port=59759,
                    db='breezy'
                )

    curs = conn.cursor()
    sql = "select * from user where username = " + "'"+str(userform)+"'"+";"
    # id 정보를 추가로 넘겨줘야 함
    print(sql)

    curs.execute(sql)
    row = curs.fetchall()
    print(row)
    auth_bridge = row[0][4]
    if auth_bridge:
        auth_list = auth_bridge.split(',')

    if (auth_list[0] == 'admin'):
        Success = 1
        pass
    else:
        for brid in auth_list:   
            if brid in BRID:
                Success = 1
                continue

    if Success == 0:
        raise Http404('교량 권한이 없습니다')
    datalist = do_test(BRID)

    time_points = list(datalist.get_points(measurement='JANET_1'))
    if (not time_points):
        raise Http404('업로드된 데이터가 없습니다')

    print(time_points[0], "ddd")
    time_list = []
    acc1_list = []
    acc2_list = []
    acc3_list = []
    acc1_sum = 0
    acc2_sum = 0
    acc3_sum = 0
    acc1_max = -999999999
    acc2_max = -999999999
    acc3_max = -999999999

    str1_list = []
    str2_list = []
    str3_list = []
    str1_sum = 0
    str2_sum = 0
    str3_sum = 0
    str1_max = -999999999
    str2_max = -999999999
    str3_max = -999999999

    cnt = 0    # 0으로 수정 필요 
    
    _list = []

    # print(time_points)
    for item in time_points:

        

        time_list.append(item["time"])
        acc1_list.append(float(item["acc_1"]))
        acc2_list.append(float(item["acc_2"]))
        acc3_list.append(float(item["acc_3"]))
        acc1_sum += float(item["acc_1"])
        acc2_sum += float(item["acc_2"])
        acc3_sum += float(item["acc_3"])
        cnt+=1
        # print(cnt)
        if float(item["acc_1"]) > acc1_max:
            acc1_max = float(item["acc_1"])
        if float(item["acc_2"]) > acc2_max:
            acc2_max = float(item["acc_2"])
        if float(item["acc_3"]) > acc3_max:
            acc3_max = float(item["acc_3"])

        str1_list.append(float(item["str_1"]))
        str2_list.append(float(item["str_2"]))
        str3_list.append(float(item["str_3"]))
        str1_sum += float(item["str_1"])
        str2_sum += float(item["str_2"])
        str3_sum += float(item["str_3"])
        if float(item["str_1"]) > str1_max:
            str1_max = float(item["str_1"])
        if float(item["str_2"]) > str2_max:
            str2_max = float(item["str_2"])
        if float(item["str_3"]) > str3_max:
            str3_max = float(item["str_3"])

        if (cnt == 60000):
            acc1 = acc1_sum/cnt   # 평균 구하기
            acc2 = acc2_sum/cnt
            acc3 = acc3_sum/cnt

            str1 = str1_sum/cnt   # 평균 구하기
            str2 = str2_sum/cnt
            str3 = str3_sum/cnt

            cnt = 0

            _list.append([time_list, acc1_list, acc2_list, acc3_list, acc1, acc2, acc3, acc1_max, acc2_max, acc3_max, str1_list, str2_list, str3_list, str1, str2, str3, str1_max, str2_max, str3_max])

            time_list = []
            acc1_list = []
            acc2_list = []
            acc3_list = []
            acc1_sum = 0
            acc2_sum = 0
            acc3_sum = 0
            acc1_max = -999999999
            acc2_max = -999999999
            acc3_max = -999999999

            str1_list = []
            str2_list = []
            str3_list = []
            str1_sum = 0
            str2_sum = 0
            str3_sum = 0
            str1_max = -999999999
            str2_max = -999999999
            str3_max = -999999999


            print(len(_list))


    
    f = open('chart/spot.txt', 'w')
    f.close()

    # f = open('chart/ddd.txt', 'w')
    # stringgg = str(_list[0][3])
    # f.write(stringgg)
    # f.close()
    for onelist in _list:
        makefgraph(onelist[3], onelist[0])
        
    
    data = np.loadtxt('chart/accgraph.txt')


    acclist = []
    spotlist = []
    doclist = []
    thread1 = Thread(target=totalgraph, args= (acclist, spotlist, BRID,))
    # totalgraph(acclist,spotlist)
    thread1.start()
    thread1.join()

    print(spotlist)
    for gjffod in spotlist:
        doclist.append(list(map(str, gjffod)))
    print(doclist)

    makereport("chart/reportform.docx", doclist, BRID)
        # print(item["time"])

        ### 리스트 형태 : _list = [[1.10분 정보][2. 10준 정보][3. 10분 정보] ... ]
        
    


    # _list = [time_list, acc1_list, acc2_list, acc3_list, acc1, acc2, acc3, acc1_max, acc2_max, acc3_max, str1_list, str2_list, str3_list, str1, str2, str3, str1_max, str2_max, str3_max]
    # reportlist = [["acc1", str(acc1), str(acc1_max)],
    #             ["acc2", str(acc2), str(acc2_max)],
    #             ["acc3", str(acc3), str(acc3_max)],
    #             ["str1", str(str1), str(str1_max)],
    #             ["str2", str(str2), str(str2_max)],
    #             ["str3", str(str3), str(str3_max)]]

    # makereport("chart/reportform.docx", reportlist)

    return render(request, 'datedata.html', {'data': data,'month': month,  \
        'year': year, 'day': day, 'month2': month2, 'year2': year2, 'day2': day2, 'list':_list[0], 'acclist':acclist, 'spotlist':spotlist, 'BRID':BRID})


def chart_data(request, cookie):
    # print(cookie)
    conn = pymysql.connect(
                    user='root',
                    passwd='pass',
                    host='43.200.68.104',
                    port=55261,
                    db='breezy'
                )
        
    curs = conn.cursor()
    sql = "select * from Sbridge_data where cookie = '"+ cookie +"';"
    curs.execute(sql)

    row = curs.fetchall()
    # print(row[0][0])  # 데이터 한 줄  
    acc1 = eval(row[0][3])
    acc2 = eval(row[0][4])

    return render(request, 'chart_data.html', {'acc1': acc1 ,'acc2': acc2})

def loading(requset, BRID, year, month, day, year2, month2, day2):
    return render(requset, 'loading.html')

def report(request, BRID):
    all_spots = []
    f = open('C:/Users/heise/Documents/janet_web/mysite/static/차트만 만드는 폴더/chart/'+BRID+'/spot.txt', 'r')
    for line in f.readlines():
            spot = line.split(' ')
            all_spots.append(spot)
    # print(spots)

    page = request.GET.get('p', 1)
    # p라는 값으로 받고, 없으면 1번째 페이지
    
    paginator = Paginator(all_spots, 10)
    # 한페이지에 2개씩
    # int를 감싸니 오류가 떠서 바꿈
    
    spots = paginator.get_page(page)

    
    return render(request, 'chartreport.html', {'BRID': BRID, 'spotlist':spots})