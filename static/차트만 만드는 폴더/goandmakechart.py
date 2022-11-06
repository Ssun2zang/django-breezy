# -*- coding: utf-8 -*-

import pymysql
from datetime import datetime, timedelta
import pprint
from influxdb import InfluxDBClient
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import numpy as np
from scipy import signal
import math
import matplotlib
matplotlib.use('Qt5agg')
import matplotlib.pyplot as plt
from threading import Thread
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches
import os
from time import sleep


def creatFolder(dir):
    if not os.path.exists(dir):
        os.makedirs(dir)

def get_ifdb(db, host='localhost', port=8086):
    
    client = InfluxDBClient(host, port, database = db)
    try:
        client.create_database(db)
        print('Connection Successful')
        # print(f'host :{host}\nport :{port}\n\ndatabase :{db}')
    except:
        print('Connection Failed')
        pass
    return client

def printresult(ifdb, tag):
    result = ifdb.query('SELECT "acc_1","acc_2","acc_3", "str_1", "str_2", "str_3" \
        FROM JANET_1 WHERE "S_id" = \''+tag+'\'')
    return result
    # pprint.pprint(result.raw)
    
def do_test(tag):  # tag -> list
    mydb = get_ifdb(db='sensor')
    # print(mydb.get_list_database())
    _list = printresult(mydb, tag)
    return _list

# Create your views here.

def makereport(_document, _list, BRID):
    document = Document(_document) 
    picname = "chart/"+BRID+ '/totalpicture.png'
    document.add_picture(picname,width= Cm(16), height= Cm(9))

    table = document.tables[0]  
    n_add_rows = len(_list) + 1 - len(table.rows) 


    for i in range(n_add_rows):  # 행 추가
        table.add_row()
    for i in range(n_add_rows+1):  # result 값을 표에 입력
        for j in range(4):
            table.cell(i+1,j).text = _list[i][j]
            table.rows[i+1].cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
    report_name = "chart/"+BRID+ "/output.docx"
    document.save(report_name) 


def makefgraph(acczlist, time, BRID):    
    b = np.array(acczlist)
    b = np.reshape(b, (len(b),1))
    fs=100
    f, Pacc = signal.csd(np.transpose(b),np.transpose(b), fs, nperseg=2048,noverlap=1024)
    f=f.reshape(-1,1)
    
    Pacc=np.transpose(Pacc)
    
    stamp = time[int(str(np.where(Pacc == max(Pacc))[0])[1:-1])][:19]
    new_stamp = stamp.replace('T', '_')


    print(math.sqrt(float(abs(max(Pacc)))))  # spot 1

    spot1 = float(f[np.where(Pacc == max(Pacc))])  
    spot2 = math.sqrt(float(abs(max(Pacc))))  
    

    f = open("chart/"+BRID+ '/spot.txt', 'a')
    stringgg = new_stamp+ " " + str(round(spot1,2)) + " " + str(round(spot2,2)) + "\n"
    f.write(stringgg)
    print(stringgg)
    f.close()



def totalgraph(_list1, _list2, BRID):
    from matplotlib import font_manager, rc
    font_path = "C:/Windows/Fonts/NGULIM.TTF"
    font = font_manager.FontProperties(fname=font_path).get_name()
    rc('font', family=font)

    data = np.loadtxt('chart/accgraph.txt')

    x=data[:,0]
    a1=data[:,1]
    a2=data[:,2]
    a3=data[:,3]


    # print(x)
    spots = []
    f = open('chart/'+BRID+'/spot.txt', 'r')
    for line in f.readlines():
            spot = line.split(' ')
            spots.append(spot)
    f.close()

    for i in range(0,len(spots)):
        t = spots[i][0]
        a = float(spots[i][1])

        b = float(spots[i][2][:-1])   

        _a = round(a, 1)   
        compa = np.where(x == _a)[0]


        result = "ERR"
        if b >= a3[compa]:
            result = "D"
        elif b >= a2[compa]:
            result = "C"
        elif b >= a1[compa]:
            result = "B"
        else:
            result = "A"
        if (b <= 0.01):
            b = 0.01

        _list2.append([t, a, b, result])

        plt.scatter(a,b,c="red", s = 20)

    f = open("chart/"+BRID+ '/spot.txt', 'w')
    for stringg in _list2:
        stringgg = "{0} {1} {2} {3} \n".format(stringg[0], stringg[1], stringg[2], stringg[3])
        f.write(stringgg)
    f.close()

    x=data[:,0].tolist()
    a1=data[:,1].tolist()
    a2=data[:,2].tolist()
    a3=data[:,3].tolist()
    _list1.append(x)
    _list1.append(a1)
    _list1.append(a2)
    _list1.append(a3)

    plt.plot(x, a1, label = 'A', color='yellowgreen')
    plt.plot(x, a2, label = 'B', color='darkorange')
    plt.plot(x, a3, label = 'C', color='darkred')
    plt.text(4.9, 0.1, 'A', fontsize=20, color='yellowgreen')
    plt.text(2,  0.55, 'B', fontsize=20, color='gold')
    plt.text(1.55, 1.4, 'C', fontsize=20, color='orange')
    plt.text(0.7, 10, 'D', fontsize=20, color='red')

    plt.xlabel('주파수(Hz)', fontsize=15, weight='bold')
    plt.ylabel('스펙트럴 밀도 (mg/√Hz)', fontsize=15, weight='bold')
    plt.semilogx()
    plt.semilogy()

    picname = "chart/"+BRID+ "/totalpicture.png"

    plt.savefig(picname)
    # plt.show()
    return




def datedata(BRID):

    datalist = do_test(BRID)  

    time_points = list(datalist.get_points(measurement='JANET_1'))
    if (not time_points):
        print("오류")

    print(time_points[0], "ddd")    # 타임포인트
    time_list = []
    acc1_list = []
    acc2_list = []
    acc3_list = []
    acc1_sum = 0
    acc2_sum = 0
    acc3_sum = 0
    acc1_max = -999999999
    acc2_max = -999999999
    acc3_max = -999999999

    str1_list = []
    str2_list = []
    str3_list = []
    str1_sum = 0
    str2_sum = 0
    str3_sum = 0
    str1_max = -999999999
    str2_max = -999999999
    str3_max = -999999999

    cnt = 0 
    
    _list = []

    # print(time_points)
    for item in time_points:

        

        time_list.append(item["time"])
        acc1_list.append(float(item["acc_1"]))
        acc2_list.append(float(item["acc_2"]))
        acc3_list.append(float(item["acc_3"]))
        acc1_sum += float(item["acc_1"])
        acc2_sum += float(item["acc_2"])
        acc3_sum += float(item["acc_3"])
        cnt+=1
        # print(cnt)
        if float(item["acc_1"]) > acc1_max:
            acc1_max = float(item["acc_1"])
        if float(item["acc_2"]) > acc2_max:
            acc2_max = float(item["acc_2"])
        if float(item["acc_3"]) > acc3_max:
            acc3_max = float(item["acc_3"])

        str1_list.append(float(item["str_1"]))
        str2_list.append(float(item["str_2"]))
        str3_list.append(float(item["str_3"]))
        str1_sum += float(item["str_1"])
        str2_sum += float(item["str_2"])
        str3_sum += float(item["str_3"])
        if float(item["str_1"]) > str1_max:
            str1_max = float(item["str_1"])
        if float(item["str_2"]) > str2_max:
            str2_max = float(item["str_2"])
        if float(item["str_3"]) > str3_max:
            str3_max = float(item["str_3"])

        if (cnt == 60000):
            acc1 = acc1_sum/cnt   # 평균 구하기
            acc2 = acc2_sum/cnt
            acc3 = acc3_sum/cnt

            str1 = str1_sum/cnt   # 평균 구하기
            str2 = str2_sum/cnt
            str3 = str3_sum/cnt

            cnt = 0

            _list.append([time_list, acc1_list, acc2_list, acc3_list, acc1, acc2, acc3, acc1_max, acc2_max, acc3_max, str1_list, str2_list, str3_list, str1, str2, str3, str1_max, str2_max, str3_max])

            time_list = []
            acc1_list = []
            acc2_list = []
            acc3_list = []
            acc1_sum = 0
            acc2_sum = 0
            acc3_sum = 0
            acc1_max = -999999999
            acc2_max = -999999999
            acc3_max = -999999999

            str1_list = []
            str2_list = []
            str3_list = []
            str1_sum = 0
            str2_sum = 0
            str3_sum = 0
            str1_max = -999999999
            str2_max = -999999999
            str3_max = -999999999


            print(len(_list))


    
    f = open('chart/'+ BRID+'/spot.txt', 'w')
    f.close()

    for onelist in _list:
        makefgraph(onelist[3], onelist[0], BRID)
        
    
    data = np.loadtxt('chart/accgraph.txt')


    acclist = []
    spotlist = []
    doclist = []


    totalgraph(acclist, spotlist, BRID)
    print(spotlist)
    for gjffod in spotlist:
        doclist.append(list(map(str, gjffod)))
    print(doclist)

    makereport("reportform.docx", doclist, BRID)

    return


def main():

    while(True):
        conn = pymysql.connect(
                user='root',
                passwd='0208',
                host='localhost',
                port=3306,
                db='breezy'
            )

        sleep(1) 
        curs = conn.cursor()
        sql = "select * from report where cur = " + str(0) + ";"
        curs.execute(sql)
        row = curs.fetchall() 
        if (row):
            creatFolder("chart/"+row[0][3])
            datedata(row[0][3])
            curs2 = conn.cursor()
            sql2 = "update report set cur = "+str(1)+" where id = "+str(row[0][0])+";"
            curs2.execute(sql2)
            conn.commit()
        print("dd")

if __name__ == "__main__":
	main()